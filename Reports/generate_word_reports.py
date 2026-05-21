from docx import Document

document = Document()

document.add_heading(
    'Hospital Management System Report',
    level=1
)

document.add_paragraph(
    'This report contains hospital patient analysis.'
)

document.add_heading('Project Features', level=2)

features = [
    "Database Connectivity",
    "Bulk Record Insertion",
    "CRUD Operations",
    "Data Cleaning",
    "CSV Export"
]

for feature in features:
    document.add_paragraph(feature)

document.save("Hospital_Report.docx")

print("Word report created")